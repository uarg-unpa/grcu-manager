#!/usr/bin/env python
"""
Script para convertir MANUAL_INSTALACION.md a formato DOCX
Genera un documento Word profesional con formato y estilos.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import re

def add_heading(doc, text, level=1):
    """Agrega un encabezado con formato personalizado"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
        heading.runs[0].font.size = Pt(24)
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(52, 73, 94)
        heading.runs[0].font.size = Pt(18)
    elif level == 3:
        heading.runs[0].font.color.rgb = RGBColor(52, 73, 94)
        heading.runs[0].font.size = Pt(14)
    return heading

def add_paragraph_with_formatting(doc, text):
    """Agrega un párrafo procesando markdown básico"""
    # Detectar si es una línea de código (comienza con espacio)
    if text.startswith('    ') or text.startswith('\t'):
        p = doc.add_paragraph(text.strip(), style='Code')
        return p
    
    # Detectar bullet points
    if text.startswith('- ') or text.startswith('* '):
        text = text[2:]
        p = doc.add_paragraph(style='List Bullet')
    # Detectar listas numeradas
    elif re.match(r'^\d+\.', text):
        text = re.sub(r'^\d+\.\s*', '', text)
        p = doc.add_paragraph(style='List Number')
    else:
        p = doc.add_paragraph()
    
    # Procesar texto en negrita y código inline (sin emojis)
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Negrita
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # Código inline
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78)
        else:
            # Texto normal
            run = p.add_run(part)
    
    return p

def add_code_block(doc, code_lines, language=''):
    """Agrega un bloque de código con formato"""
    if language:
        doc.add_paragraph(f'[{language}]', style='Intense Quote')
    
    for line in code_lines:
        p = doc.add_paragraph(line, style='Code')
        p.paragraph_format.left_indent = Inches(0.5)

def add_cover_page(doc):
    """Agrega una carátula profesional al documento"""
    # Título principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Universidad Nacional de la Patagonia Austral')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    
    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Laboratorio de Desarrollo de Software')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(52, 73, 94)
    
    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio
    
    # Título del manual
    manual_title = doc.add_paragraph()
    manual_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = manual_title.add_run('MANUAL DE INSTALACIÓN')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(41, 128, 185)
    
    # Subtítulo del sistema
    system_title = doc.add_paragraph()
    system_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = system_title.add_run('GRCU Manager')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(52, 73, 94)
    
    system_subtitle = doc.add_paragraph()
    system_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = system_subtitle.add_run('Sistema de Gestión de Requerimientos y Casos de Uso')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio
    
    # Grupo
    group_title = doc.add_paragraph()
    group_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = group_title.add_run('Grupo 4Bytes')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    
    doc.add_paragraph()  # Espacio
    
    # Integrantes
    integrantes_title = doc.add_paragraph()
    integrantes_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = integrantes_title.add_run('Integrantes:')
    run.font.size = Pt(12)
    run.font.bold = True
    
    integrantes = [
        'Nicolás Butterfield',
        'Martina Gagna',
        'Abril Alvarez',
        'Cristian Carranza'
    ]
    
    for integrante in integrantes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(integrante)
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio
    
    # Docentes
    docentes_title = doc.add_paragraph()
    docentes_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = docentes_title.add_run('Docentes:')
    run.font.size = Pt(12)
    run.font.bold = True
    
    docentes = [
        'Lic. Karin Hallar',
        'Lic. Esteban Gesto',
        'Mg. Osiris Sofía'
    ]
    
    for docente in docentes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(docente)
        run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espacio
    doc.add_paragraph()  # Espacio
    
    # Fecha
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fecha.add_run('Noviembre 2025')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(127, 140, 141)
    
    # Salto de página después de la carátula
    doc.add_page_break()

def add_table_of_contents(doc):
    """Agrega una página de índice"""
    # Título del índice
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('ÍNDICE')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    
    doc.add_paragraph()  # Espacio
    
    # Contenido del índice (manual)
    toc_items = [
        ('1. Introducción', ''),
        ('2. Requisitos del Sistema', ''),
        ('   • Software Requerido', ''),
        ('   • Requisitos de Hardware Mínimos', ''),
        ('   • Sistemas Operativos Soportados', ''),
        ('3. Instalación para Desarrollo', ''),
        ('   • Paso 1: Instalación de Python', ''),
        ('   • Paso 2: Clonar el Repositorio', ''),
        ('   • Paso 3: Crear Entorno Virtual', ''),
        ('   • Paso 4: Instalar Dependencias', ''),
        ('   • Paso 5: Configurar Variables de Entorno', ''),
        ('   • Paso 6: Configurar Base de Datos', ''),
        ('   • Paso 7: Aplicar Migraciones', ''),
        ('   • Paso 8: Crear Directorios de Medios', ''),
        ('   • Paso 9: Ejecutar el Servidor de Desarrollo', ''),
        ('   • Paso 10: Acceder a la Aplicación', ''),
        ('4. Instalación para Producción', ''),
        ('   • Consideraciones de Producción', ''),
        ('   • Configuración Básica de Producción', ''),
        ('5. Configuración Inicial', ''),
        ('   • Primera Ejecución', ''),
        ('   • Configuración de Google OAuth', ''),
        ('   • Cargar Datos de Demostración', ''),
        ('6. Solución de Problemas', ''),
        ('   • Errores Comunes y Soluciones', ''),
        ('7. Recursos Adicionales', ''),
        ('   • Documentación Oficial', ''),
        ('   • Archivos de Configuración Importantes', ''),
        ('   • Comandos Útiles de Django', ''),
        ('8. Soporte', ''),
        ('9. Licencia', ''),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        if item.startswith('   •'):
            # Subítem
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(item.strip())
            run.font.size = Pt(10)
        elif item.startswith('   '):
            # Subítem con bullet
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(item.strip())
            run.font.size = Pt(10)
        else:
            # Ítem principal
            run = p.add_run(item)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(52, 73, 94)
    
    # Salto de página después del índice
    doc.add_page_break()

def convert_markdown_to_docx(md_file, docx_file):
    """Convierte archivo Markdown a DOCX"""
    
    # Crear documento
    doc = Document()
    
    # Agregar carátula
    add_cover_page(doc)
    
    # Agregar índice
    add_table_of_contents(doc)
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo para código
    try:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(9)
        code_font.color.rgb = RGBColor(0, 0, 0)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Inches(0.25)
    except:
        code_style = styles['Normal']
    
    # Leer archivo markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    code_block_lines = []
    code_language = ''
    skip_next = False
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        if skip_next:
            skip_next = False
            i += 1
            continue
        
        # Detectar bloques de código
        if line.startswith('```'):
            if in_code_block:
                # Fin del bloque de código
                add_code_block(doc, code_block_lines, code_language)
                code_block_lines = []
                code_language = ''
                in_code_block = False
            else:
                # Inicio del bloque de código
                code_language = line[3:].strip()
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Omitir líneas vacías múltiples
        if not line:
            doc.add_paragraph()
            i += 1
            continue
        
        # Detectar encabezados
        if line.startswith('# '):
            add_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=4)
        # Omitir líneas de separación
        elif line.strip() == '---':
            doc.add_paragraph()
        # Omitir HTML
        elif line.strip().startswith('<') or line.strip().endswith('>'):
            pass
        else:
            add_paragraph_with_formatting(doc, line)
        
        i += 1
    
    # Configurar márgenes y tamaño de página, encabezado y pie de página
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Encabezado con logo (izquierda) y título (derecha)
        header = section.header
        header.is_linked_to_previous = False
        # Crear una tabla de 2 columnas para logo + texto
        header_table = header.add_table(rows=1, cols=2, width=Inches(6))
        header_table.columns[0].width = Inches(2.5)
        header_table.columns[1].width = Inches(3.5)
        header_table.autofit = True
        try:
            logo_path = 'accounts/static/accounts/img/logo_grcu_manager.png'
            left_cell = header_table.cell(0, 0)
            p_left = left_cell.paragraphs[0]
            run_left = p_left.add_run()
            run_left.add_picture(logo_path, height=Inches(0.5))
        except Exception:
            # Si falla el logo, deja la celda vacía
            pass

        right_cell = header_table.cell(0, 1)
        p_right = right_cell.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_right = p_right.add_run("Manual de Instalación - GRCU Manager")
        run_right.font.size = Pt(10)
        run_right.font.bold = True
        run_right.font.color.rgb = RGBColor(44, 62, 80)

        # Pie de página con número de página y fecha de generación
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Texto izquierdo
        run_footer_text = footer_paragraph.add_run("Universidad Nacional de la Patagonia Austral • GRCU Manager • Página ")
        run_footer_text.font.size = Pt(10)
        run_footer_text.font.color.rgb = RGBColor(127, 140, 141)

        # Campo PAGE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        page_run = footer_paragraph.add_run()
        page_run.font.size = Pt(10)
        page_run.font.color.rgb = RGBColor(127, 140, 141)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        page_run._r.append(fldChar1)
        page_run._r.append(instrText)
        page_run._r.append(fldChar2)

        # Texto ' de '
        run_de = footer_paragraph.add_run(' de ')
        run_de.font.size = Pt(10)
        run_de.font.color.rgb = RGBColor(127, 140, 141)

        # Campo NUMPAGES
        numpages_run = footer_paragraph.add_run()
        numpages_run.font.size = Pt(10)
        numpages_run.font.color.rgb = RGBColor(127, 140, 141)
        fldChar1b = OxmlElement('w:fldChar')
        fldChar1b.set(qn('w:fldCharType'), 'begin')
        instrTextB = OxmlElement('w:instrText')
        instrTextB.text = 'NUMPAGES'
        fldChar2b = OxmlElement('w:fldChar')
        fldChar2b.set(qn('w:fldCharType'), 'end')
        numpages_run._r.append(fldChar1b)
        numpages_run._r.append(instrTextB)
        numpages_run._r.append(fldChar2b)

        # Fecha de generación
        run_date_sep = footer_paragraph.add_run(' • Generado: ')
        run_date_sep.font.size = Pt(10)
        run_date_sep.font.color.rgb = RGBColor(127, 140, 141)
        run_date = footer_paragraph.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
        run_date.font.size = Pt(10)
        run_date.font.color.rgb = RGBColor(127, 140, 141)

    # Guardar documento
    doc.save(docx_file)
    print(f'✅ Documento generado exitosamente: {docx_file}')

if __name__ == '__main__':
    input_file = 'MANUAL_INSTALACION.md'
    output_file = 'MANUAL_INSTALACION.docx'
    
    try:
        convert_markdown_to_docx(input_file, output_file)
        print(f'\n📄 Manual convertido a formato Word')
        print(f'📁 Ubicación: {output_file}')
    except Exception as e:
        print(f'❌ Error al generar el documento: {e}')
        import traceback
        traceback.print_exc()
