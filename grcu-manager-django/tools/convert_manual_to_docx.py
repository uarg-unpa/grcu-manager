#!/usr/bin/env python3
"""Convertir MANUAL_INSTALACION.md a MANUAL_INSTALACION.docx

Este script hace una conversión sencilla (encabezados, párrafos, bloques de código)
e inserta el logo del proyecto en el header y números de página en el footer.

Uso:
    python tools/convert_manual_to_docx.py

Requiere: python-docx, Pillow
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / 'MANUAL_INSTALACION.md'
DOCX_PATH = ROOT / 'MANUAL_INSTALACION.docx'
LOGO_PATH = ROOT / 'accounts' / 'static' / 'accounts' / 'img' / 'favicon.png'
LOGO_4BYTES_PATH = ROOT / 'accounts' / 'static' / 'accounts' / 'img' / 'logo_4bytes.jpg'


def add_page_number(paragraph):
    """Insert a PAGE field into a paragraph (shows current page number)."""
    run = paragraph.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)


def add_num_pages(paragraph):
    run = paragraph.add_run(' / ')
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'NUMPAGES')
    run._r.append(fld)


def make_doc():
    if not MD_PATH.exists():
        print('No existe', MD_PATH)
        return 1

    doc = Document()

    # core properties / metadata
    try:
        doc.core_properties.title = 'Manual de instalación — GRCU Manager'
        doc.core_properties.author = 'Nicolás Butterfield; Abril Alvarez; Martina Gagna; Cristian Carranza'
        doc.core_properties.comments = 'Contacto: nicbutter@gmail.com'
    except Exception:
        pass

    # Improve base styles and margins
    sections = doc.sections
    for sec in sections:
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

    # Set default font for Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    # Heading styles
    try:
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(18)
        h1.font.bold = True
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(14)
        h2.font.bold = True
    except Exception:
        pass

    # Cover page: logo big + title + date
    from datetime import date as _date
    cover_section = doc.sections[0]
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if LOGO_PATH.exists():
        try:
            run = cover_para.add_run()
            run.add_picture(str(LOGO_PATH), width=Cm(6))
        except Exception as e:
            print('No se pudo insertar logo en portada:', e)

    title = 'Manual de instalación — GRCU Manager'
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p_title.add_run(title)
    r.bold = True
    r.font.size = Pt(24)

    subtitle = f"Última actualización: {_date.today().isoformat()}"
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    rs = p_sub.add_run(subtitle)
    rs.italic = True
    rs.font.size = Pt(10)

    # Authors / Copyright page
    doc.add_page_break()
    a_para = doc.add_paragraph()
    a_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    a_title = a_para.add_run('Autores y derechos')
    a_title.bold = True
    a_title.font.size = Pt(14)


    doc.add_paragraph()
    authors = doc.add_paragraph()
    authors.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    authors.add_run('Grupo 4Bytes').bold = True
    doc.add_paragraph()
    miembros = doc.add_paragraph()
    miembros.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    miembros.add_run('Nicolás Butterfield, Abril Alvarez, Martina Gagna, Cristian Carranza')
    doc.add_paragraph()
    uni = doc.add_paragraph()
    uni.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    uni.add_run('Universidad Nacional de la Patagonia Austral').bold = False

    doc.add_paragraph()
    copy = doc.add_paragraph()
    copy.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    copy.add_run('\u00A9 Grupo 4Bytes & Universidad Nacional de la Patagonia Austral. Todos los derechos reservados.').italic = False

    doc.add_page_break()

    # header with logos and footer with page numbers for all sections
    for section in doc.sections:
        header = section.header
        try:
            tbl = header.add_table(rows=1, cols=2)
            tbl.autofit = False
            tbl.columns[0].width = Cm(9)
            tbl.columns[1].width = Cm(9)
            row = tbl.rows[0]
            cell_left = row.cells[0]
            cell_right = row.cells[1]
            p_left = cell_left.paragraphs[0]
            p_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            if LOGO_PATH.exists():
                try:
                    run = p_left.add_run()
                    run.add_picture(str(LOGO_PATH), width=Cm(2))
                except Exception:
                    p_left.add_run('GRCU Manager')
            else:
                p_left.add_run('GRCU Manager')

            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            if LOGO_4BYTES_PATH.exists():
                try:
                    runr = p_right.add_run()
                    runr.add_picture(str(LOGO_4BYTES_PATH), width=Cm(3))
                except Exception:
                    p_right.add_run('Grupo 4Bytes')
            else:
                p_right.add_run('Grupo 4Bytes')
        except Exception:
            # fallback header
            header.paragraphs[0].add_run('GRCU Manager - Grupo 4Bytes')

        # footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fp.clear()
        fp.add_run('Página ')
        add_page_number(fp)
        add_num_pages(fp)

    in_code = False
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            if line.strip().startswith('```'):
                in_code = not in_code
                if in_code:
                    # start code block
                    p = doc.add_paragraph()
                    p.style.font.name = 'Courier New'
                    p.style.font.size = Pt(9)
                else:
                    # end code block
                    p = None
                continue

            if in_code:
                # append to last paragraph
                if doc.paragraphs:
                    doc.paragraphs[-1].add_run(line + '\n')
                else:
                    doc.add_paragraph(line)
                continue

            # headings
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.strip() == '---' or line.strip() == '***':
                doc.add_page_break()
            elif line.strip().startswith('```'):
                # handled above
                continue
            elif line.strip() == '':
                doc.add_paragraph('')
            else:
                # detect code fence inlined or preformatted blocks marked with 4 spaces
                if line.startswith('    '):
                    p = doc.add_paragraph()
                    p.style.font.name = 'Courier New'
                    p.style.font.size = Pt(9)
                    p.add_run(line[4:])
                else:
                    doc.add_paragraph(line)

    doc.save(DOCX_PATH)
    print('Generado:', DOCX_PATH)
    return 0


if __name__ == '__main__':
    raise SystemExit(make_doc())
